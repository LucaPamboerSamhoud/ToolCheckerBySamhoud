import io
from datetime import datetime

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

from ..models import ComplianceResult, TrafficLight

# &samhoud kleuren
SAMHOUD_BLUE = RGBColor(0x0C, 0x2A, 0xAD)
SAMHOUD_LIGHT_BLUE = RGBColor(0x42, 0x72, 0xAB)
SAMHOUD_BLUE_SOFT = RGBColor(0x70, 0x8D, 0xB8)
SAMHOUD_ORANGE = RGBColor(0xFF, 0xA3, 0x14)
COLOR_GREEN = RGBColor(0x16, 0xA3, 0x4A)
COLOR_ORANGE = RGBColor(0xEA, 0x58, 0x0C)
COLOR_RED = RGBColor(0xDC, 0x26, 0x26)
COLOR_GRAY = RGBColor(0x64, 0x74, 0x8B)
COLOR_DARK = RGBColor(0x1A, 0x1A, 0x2E)
COLOR_WHITE = RGBColor(0xFF, 0xFF, 0xFF)

STATUS_COLORS = {
    TrafficLight.GREEN: COLOR_GREEN,
    TrafficLight.ORANGE: COLOR_ORANGE,
    TrafficLight.RED: COLOR_RED,
}

STATUS_LABELS = {
    TrafficLight.GREEN: "GROEN — Veilig te gebruiken",
    TrafficLight.ORANGE: "ORANJE — Nader onderzoek nodig",
    TrafficLight.RED: "ROOD — Niet gebruiken zonder waarborgen",
}

STATUS_EMOJI = {
    TrafficLight.GREEN: "V",
    TrafficLight.ORANGE: "!",
    TrafficLight.RED: "X",
}

# Educatieve content per categorie
CATEGORY_EDUCATION = {
    "Dataopslag & Verwerking": (
        "Waar je data staat bepaalt welke wetten van toepassing zijn. "
        "Data buiten de EU valt niet automatisch onder de bescherming van de AVG. "
        "Dat betekent dat persoonsgegevens mogelijk minder goed beschermd zijn. "
        "Ook sub-verwerkers (bedrijven die namens de tool data verwerken) spelen een rol: "
        "als een sub-verwerker data in de VS verwerkt, is dat een risicofactor."
    ),
    "Datarechten (AVG)": (
        "De AVG geeft iedereen rechten over hun eigen data: het recht op inzage, "
        "correctie, verwijdering (recht op vergetelheid), en dataportabiliteit. "
        "Als een tool deze rechten niet ondersteunt, is het juridisch risicovol "
        "om er persoonsgegevens in te verwerken. Let ook op of data gebruikt "
        "wordt voor het trainen van AI-modellen — dit vereist expliciete toestemming."
    ),
    "Beveiliging": (
        "Goede beveiliging is de basis van privacy. Zonder encryptie, certificeringen "
        "en een plan voor datalekken zijn persoonsgegevens kwetsbaar. Een datalek kan "
        "leiden tot boetes tot 4% van de jaaromzet. Zoek naar ISO 27001 en SOC 2 "
        "certificeringen — dit zijn internationale standaarden die aantonen dat een "
        "organisatie beveiliging serieus neemt."
    ),
}


def generate_report(result: ComplianceResult) -> io.BytesIO:
    """Genereer een Word rapport voor een compliance check resultaat."""
    doc = Document()
    _set_default_font(doc)

    # Cover pagina
    _add_cover_page(doc, result)

    # Managementsamenvatting
    _add_summary(doc, result)

    # Categorieën met checks
    for category in result.categories:
        _add_category(doc, category)

    # Sub-verwerkers
    if result.sub_processors:
        _add_sub_processors(doc, result)

    # Aanbevelingen
    _add_recommendations(doc, result)

    # Geraadpleegde bronnen
    _add_sources(doc, result)

    # Disclaimer
    _add_disclaimer(doc, result)

    # Footer info
    _add_footer_info(doc)

    # Sla op in bytes buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _set_default_font(doc: Document) -> None:
    """Stel het standaard font in."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)
    font.color.rgb = COLOR_DARK


def _add_cover_page(doc: Document, result: ComplianceResult) -> None:
    """Voeg een cover pagina toe."""
    # Witruimte bovenaan
    for _ in range(6):
        doc.add_paragraph()

    # &samhoud label
    label = doc.add_paragraph()
    label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = label.add_run("&samhoud  |  ToolChecker")
    run.font.color.rgb = SAMHOUD_BLUE
    run.font.size = Pt(11)

    # Hoofdtitel
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.space_after = Pt(4)
    run = title.add_run(f"AVG/GDPR Compliance Rapport")
    run.font.color.rgb = SAMHOUD_BLUE
    run.font.size = Pt(28)
    run.bold = True

    # Tool naam
    tool = doc.add_paragraph()
    tool.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tool.space_after = Pt(20)
    run = tool.add_run(result.tool_name)
    run.font.color.rgb = COLOR_DARK
    run.font.size = Pt(22)

    # Status indicator
    status_para = doc.add_paragraph()
    status_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    status_color = STATUS_COLORS.get(result.overall_status, COLOR_GRAY)
    status_label = STATUS_LABELS.get(result.overall_status, "Onbekend")

    indicator = status_para.add_run(f"  {STATUS_EMOJI[result.overall_status]}  ")
    indicator.font.size = Pt(18)
    indicator.bold = True
    indicator.font.color.rgb = status_color

    label_run = status_para.add_run(status_label)
    label_run.font.size = Pt(14)
    label_run.bold = True
    label_run.font.color.rgb = status_color

    # Witruimte
    for _ in range(4):
        doc.add_paragraph()

    # Meta informatie
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run(
        f"Gegenereerd op {datetime.now().strftime('%d %B %Y om %H:%M')}\n"
        f"Door ToolChecker by &samhoud"
    )
    meta_run.font.color.rgb = COLOR_GRAY
    meta_run.font.size = Pt(9)

    if result.tool_url:
        url_para = doc.add_paragraph()
        url_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        url_run = url_para.add_run(result.tool_url)
        url_run.font.color.rgb = SAMHOUD_LIGHT_BLUE
        url_run.font.size = Pt(9)

    # Pagina-einde
    doc.add_page_break()


def _add_summary(doc: Document, result: ComplianceResult) -> None:
    """Voeg de managementsamenvatting toe."""
    heading = doc.add_heading("Managementsamenvatting", level=1)
    for run in heading.runs:
        run.font.color.rgb = SAMHOUD_BLUE

    # Uitleg wat dit rapport is
    intro = doc.add_paragraph()
    intro_run = intro.add_run(
        f"Dit rapport bevat de resultaten van een geautomatiseerde AVG/GDPR compliance "
        f"check van {result.tool_name}. De analyse is gebaseerd op publiek beschikbare "
        f"informatie zoals de privacy policy, security documentatie en sub-verwerkerlijsten. "
        f"Per categorie is een stoplicht-oordeel gegeven."
    )
    intro_run.font.size = Pt(10)

    doc.add_paragraph()

    # AI samenvatting
    doc.add_paragraph(result.summary)

    # Overzichtstabel
    doc.add_paragraph()
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0].cells
    _set_cell_text(hdr[0], "Categorie", bold=True, color=SAMHOUD_BLUE)
    _set_cell_text(hdr[1], "Beoordeling", bold=True, color=SAMHOUD_BLUE)

    for cat in result.categories:
        row = table.add_row().cells
        _set_cell_text(row[0], cat.name)
        color = STATUS_COLORS.get(cat.status, COLOR_GRAY)
        _set_cell_text(
            row[1], STATUS_LABELS.get(cat.status, "Onbekend"), color=color, bold=True
        )

    doc.add_paragraph()

    # Stoplicht uitleg
    legend = doc.add_paragraph()
    legend_run = legend.add_run("Stoplicht betekenis:")
    legend_run.bold = True
    legend_run.font.size = Pt(9)
    legend_run.font.color.rgb = COLOR_GRAY

    for status, label, desc in [
        (TrafficLight.GREEN, "GROEN", "De tool voldoet op dit punt aan de AVG-vereisten."),
        (TrafficLight.ORANGE, "ORANJE", "Er zijn aandachtspunten of er kon onvoldoende informatie gevonden worden."),
        (TrafficLight.RED, "ROOD", "Er zijn serieuze risico's gevonden die actie vereisen."),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        color = STATUS_COLORS[status]
        label_run = p.add_run(f"{label}: ")
        label_run.bold = True
        label_run.font.size = Pt(9)
        label_run.font.color.rgb = color
        desc_run = p.add_run(desc)
        desc_run.font.size = Pt(9)
        desc_run.font.color.rgb = COLOR_GRAY

    doc.add_paragraph()


def _add_category(doc: Document, category) -> None:
    """Voeg een categorie sectie toe met alle checks."""
    heading = doc.add_heading(category.name, level=2)
    for run in heading.runs:
        run.font.color.rgb = SAMHOUD_LIGHT_BLUE

    # Status
    status_para = doc.add_paragraph()
    status_color = STATUS_COLORS.get(category.status, COLOR_GRAY)
    run = status_para.add_run("Beoordeling: ")
    run.bold = True
    status_run = status_para.add_run(
        STATUS_LABELS.get(category.status, "Onbekend")
    )
    status_run.font.color.rgb = status_color
    status_run.bold = True

    doc.add_paragraph(category.summary)

    # Educatieve uitleg
    education = CATEGORY_EDUCATION.get(category.name)
    if education:
        edu_para = doc.add_paragraph()
        edu_label = edu_para.add_run("Waarom is dit belangrijk? ")
        edu_label.bold = True
        edu_label.font.size = Pt(9)
        edu_label.font.color.rgb = SAMHOUD_BLUE
        edu_text = edu_para.add_run(education)
        edu_text.font.size = Pt(9)
        edu_text.font.color.rgb = COLOR_GRAY
        edu_text.italic = True

    doc.add_paragraph()

    # Individuele checks
    for check in category.checks:
        _add_check(doc, check)


def _add_check(doc: Document, check) -> None:
    """Voeg een individuele check toe."""
    # Check naam met status indicator
    para = doc.add_paragraph()
    status_color = STATUS_COLORS.get(check.status, COLOR_GRAY)

    indicator = para.add_run(f"[{STATUS_EMOJI[check.status]}] ")
    indicator.font.color.rgb = status_color
    indicator.bold = True

    name_run = para.add_run(check.name)
    name_run.bold = True

    # Bevinding
    finding = doc.add_paragraph()
    finding.paragraph_format.left_indent = Cm(0.5)
    finding.add_run(check.finding)

    # Bronnen met quotes
    if check.sources:
        for source in check.sources:
            source_para = doc.add_paragraph()
            source_para.paragraph_format.left_indent = Cm(1)

            if source.quote:
                quote_run = source_para.add_run(f'"{source.quote}"')
                quote_run.italic = True
                quote_run.font.size = Pt(9)
                quote_run.font.color.rgb = SAMHOUD_BLUE_SOFT
                source_para.add_run("\n")

            link_run = source_para.add_run(f"Bron: {source.url}")
            link_run.font.size = Pt(8)
            link_run.font.color.rgb = SAMHOUD_LIGHT_BLUE


def _add_sub_processors(doc: Document, result: ComplianceResult) -> None:
    """Voeg het sub-verwerkers overzicht toe."""
    heading = doc.add_heading("Sub-verwerkers", level=2)
    for run in heading.runs:
        run.font.color.rgb = SAMHOUD_LIGHT_BLUE

    # Uitleg
    intro = doc.add_paragraph()
    intro_run = intro.add_run(
        f"{result.tool_name} maakt gebruik van andere bedrijven (sub-verwerkers) "
        f"om data te verwerken of op te slaan. De veiligheid van je data is zo sterk "
        f"als de zwakste schakel in deze keten. Als een sub-verwerker data buiten de EU "
        f"verwerkt, heeft dit gevolgen voor de AVG-compliance."
    )
    intro_run.font.size = Pt(10)

    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0].cells
    _set_cell_text(hdr[0], "Sub-verwerker", bold=True, color=SAMHOUD_BLUE)
    _set_cell_text(hdr[1], "Doel", bold=True, color=SAMHOUD_BLUE)
    _set_cell_text(hdr[2], "Datalocatie", bold=True, color=SAMHOUD_BLUE)
    _set_cell_text(hdr[3], "Status", bold=True, color=SAMHOUD_BLUE)

    for sp in result.sub_processors:
        row = table.add_row().cells
        _set_cell_text(row[0], sp.name)
        _set_cell_text(row[1], sp.purpose)
        _set_cell_text(row[2], sp.data_location)
        color = STATUS_COLORS.get(sp.status, COLOR_GRAY)
        _set_cell_text(row[3], sp.status.value.upper(), color=color, bold=True)

    doc.add_paragraph()


def _add_recommendations(doc: Document, result: ComplianceResult) -> None:
    """Voeg aanbevelingen toe op basis van het resultaat."""
    heading = doc.add_heading("Aanbevelingen", level=2)
    for run in heading.runs:
        run.font.color.rgb = SAMHOUD_LIGHT_BLUE

    # Generieke aanbevelingen op basis van status
    if result.overall_status == TrafficLight.GREEN:
        recs = [
            "Lees de verwerkersovereenkomst (DPA) door voordat je de tool inzet voor persoonsgegevens.",
            "Documenteer in je verwerkingsregister welke persoonsgegevens je met deze tool verwerkt.",
            "Herhaal deze check periodiek, omdat tools hun beleid kunnen wijzigen.",
        ]
    elif result.overall_status == TrafficLight.ORANGE:
        recs = [
            "Neem contact op met de leverancier om de openstaande vragen te beantwoorden.",
            "Schakel je Functionaris Gegevensbescherming (FG) in voor een definitieve beoordeling.",
            "Overweeg een Data Protection Impact Assessment (DPIA) uit te voeren.",
            "Gebruik de tool voorlopig niet voor gevoelige persoonsgegevens tot de onduidelijkheden zijn opgelost.",
            "Documenteer de risico's en genomen maatregelen in je verwerkingsregister.",
        ]
    else:
        recs = [
            "Gebruik deze tool NIET voor het verwerken van persoonsgegevens zonder aanvullende maatregelen.",
            "Schakel juridisch advies in en overleg met je Functionaris Gegevensbescherming (FG).",
            "Voer een Data Protection Impact Assessment (DPIA) uit.",
            "Onderzoek alternatieven die beter aan de AVG-vereisten voldoen.",
            "Als de tool onmisbaar is, bespreek met de leverancier welke aanpassingen mogelijk zijn.",
        ]

    for rec in recs:
        para = doc.add_paragraph(style="List Bullet")
        run = para.add_run(rec)
        run.font.size = Pt(10)

    doc.add_paragraph()

    # CTA
    cta = doc.add_paragraph()
    cta_run = cta.add_run(
        "Hulp nodig bij het implementeren van deze aanbevelingen? "
        "Het Data & AI team van &samhoud helpt je graag. "
        "Neem contact op via data.team@samhoud.com"
    )
    cta_run.font.color.rgb = SAMHOUD_BLUE
    cta_run.bold = True
    cta_run.font.size = Pt(10)

    doc.add_paragraph()


def _add_sources(doc: Document, result: ComplianceResult) -> None:
    """Voeg de lijst met geraadpleegde bronnen toe."""
    heading = doc.add_heading("Geraadpleegde bronnen", level=2)
    for run in heading.runs:
        run.font.color.rgb = SAMHOUD_LIGHT_BLUE

    intro = doc.add_paragraph()
    intro_run = intro.add_run(
        "De volgende bronnen zijn geraadpleegd tijdens deze analyse:"
    )
    intro_run.font.size = Pt(10)
    intro_run.font.color.rgb = COLOR_GRAY

    for source in result.sources_consulted:
        para = doc.add_paragraph(style="List Bullet")
        title_run = para.add_run(source.title)
        title_run.bold = True
        title_run.font.size = Pt(9)
        para.add_run("\n")
        url_run = para.add_run(source.url)
        url_run.font.size = Pt(8)
        url_run.font.color.rgb = SAMHOUD_LIGHT_BLUE


def _add_disclaimer(doc: Document, result: ComplianceResult) -> None:
    """Voeg de disclaimer toe."""
    doc.add_paragraph()
    heading = doc.add_heading("Disclaimer", level=2)
    for run in heading.runs:
        run.font.color.rgb = COLOR_GRAY

    disclaimer_para = doc.add_paragraph()
    run = disclaimer_para.add_run(result.disclaimer)
    run.font.color.rgb = COLOR_GRAY
    run.font.size = Pt(9)
    run.italic = True

    # Extra disclaimer
    extra = doc.add_paragraph()
    extra_run = extra.add_run(
        "Deze analyse is gebaseerd op publiek beschikbare informatie op het moment van de check. "
        "Tools kunnen hun beleid, documentatie en technische implementatie wijzigen. "
        "Het is aan te raden om deze check periodiek te herhalen en de uitkomsten te bespreken "
        "met een Functionaris Gegevensbescherming (FG) of juridisch adviseur."
    )
    extra_run.font.color.rgb = COLOR_GRAY
    extra_run.font.size = Pt(9)
    extra_run.italic = True


def _add_footer_info(doc: Document) -> None:
    """Voeg &samhoud footer informatie toe."""
    doc.add_paragraph()
    doc.add_paragraph()

    # Scheidingslijn simulatie
    line = doc.add_paragraph()
    line_run = line.add_run("_" * 60)
    line_run.font.color.rgb = RGBColor(0xCB, 0xD6, 0xE5)
    line_run.font.size = Pt(8)

    # &samhoud info
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER

    brand = info.add_run("&samhoud")
    brand.bold = True
    brand.font.color.rgb = SAMHOUD_BLUE
    brand.font.size = Pt(11)

    info.add_run("\n")

    tagline = info.add_run("Together we build a brighter future")
    tagline.italic = True
    tagline.font.color.rgb = SAMHOUD_LIGHT_BLUE
    tagline.font.size = Pt(9)

    info.add_run("\n\n")

    contact = info.add_run("data.team@samhoud.com  |  samhoud.com")
    contact.font.color.rgb = COLOR_GRAY
    contact.font.size = Pt(8)


def _set_cell_text(
    cell, text: str, bold: bool = False, color: RGBColor | None = None
) -> None:
    """Helper om tekst in een tabelcel te zetten."""
    cell.text = ""
    para = cell.paragraphs[0]
    run = para.add_run(text)
    run.font.size = Pt(9)
    run.bold = bold
    if color:
        run.font.color.rgb = color
